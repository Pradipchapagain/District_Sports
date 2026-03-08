import pandas as pd
import database as db
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import qrcode

# ==========================================
# 🛠️ Helper Function: सेलको ब्याकग्राउन्ड रङ परिवर्तन गर्ने
# ==========================================
def set_cell_background(cell, color_hex):
    """Word Table को सेलमा Background Color भर्ने फङ्सन"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shade = OxmlElement('w:shd')
    shade.set(qn('w:val'), 'clear')
    shade.set(qn('w:color'), 'auto')
    shade.set(qn('w:fill'), color_hex)
    tcPr.append(shade)

# ==========================================
# 📜 परिचय-पत्र बनाउने फङ्सन (Optimized & Advanced)
# ==========================================
def generate_id_cards_docx(school_name, players_df):
    """Word फाइलमा आकर्षक परिचयपत्र (ID Cards) बनाउने फङ्सन"""
    
    # ७. Validation: डाटा नभएमा रोक्ने
    if players_df is None or players_df.empty:
        return None

    # १. Bulk DB Query (N+1 समस्याको समाधान)
    player_ids = tuple(players_df['id'].tolist())
    events_dict = {}
    if player_ids:
        conn = db.get_connection()
        # PostgreSQL को लागि IN मा tuple पठाउँदा %s प्रयोग हुन्छ
        placeholders = ', '.join(['%s'] * len(player_ids)) 
        query = f"""
            SELECT r.player_id, e.name 
            FROM registrations r
            JOIN events e ON r.event_code = e.code
            WHERE r.player_id IN ({placeholders})
        """
        events_raw = pd.read_sql_query(query, conn, params=player_ids)
        conn.close()
        
        # खेलाडी अनुसार इभेन्टको सूची (Dictionary) बनाउने
        if not events_raw.empty:
            events_dict = events_raw.groupby('player_id')['name'].apply(list).to_dict()

    # डकुमेन्ट सेटअप
    doc = Document()
    
    # ८. Margin Setup (ID कार्डको लागि साँघुरो मार्जिन)
    section = doc.sections[0]
    section.page_width = Cm(21.0)  # A4 चौडाइ
    section.page_height = Cm(29.7) # A4 उचाइ
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    
    # Main Header
    head = doc.add_heading('१६औं जिल्ला स्तरीय राष्ट्रपति रनिङ शिल्ड प्रतियोगिता २०८२', level=1)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_head = head.runs[0]
    run_head.font.size = Pt(14)
    run_head.font.color.rgb = RGBColor(0, 51, 102)
    # ४. नेपाली फन्ट सपोर्ट
    run_head.font.name = 'Arial Unicode MS'
    run_head._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    
    sub = doc.add_paragraph(f"विद्यालय: {school_name}")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.name = 'Arial Unicode MS'
    
    doc.add_paragraph("") # Space
    
    # ५. Grid Table (2 Columns, Fixed Height & Width)
    rows_needed = (len(players_df) + 1) // 2
    table = doc.add_table(rows=rows_needed, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for idx, (_, player) in enumerate(players_df.iterrows()):
        r = idx // 2
        c = idx % 2
        
        row = table.rows[r]
        row.height = Cm(6.5) # ५. कार्डको फिक्स उचाइ
        
        cell = table.cell(r, c)
        cell.width = Cm(9.0) # कार्डको फिक्स चौडाइ
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # ८. हल्का नीलो ब्याकग्राउन्ड
        set_cell_background(cell, 'F8FAFC')
        
        # ३. छुट्टाछुट्टै Paragraph को प्रयोग
        # --- Title ---
        p_title = cell.paragraphs[0]
        p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r_title = p_title.add_run("PLAYER IDENTITY CARD")
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(185, 28, 28) # Red
        
        # --- Photo & QR Code (Side by Side visually) ---
        p_media = cell.add_paragraph()
        p_media.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # २. Photo Handling
        photo_path = player.get('photo_path', '')
        if pd.notna(photo_path) and os.path.exists(str(photo_path)):
            # यदि सक्कली फोटो छ भने राख्ने
            p_media.add_run().add_picture(str(photo_path), height=Cm(2.5))
        else:
            # छैन भने प्लेसहोल्डर
            r_photo = p_media.add_run("\n[ PP Size Photo ]\n")
            r_photo.font.size = Pt(9)
            r_photo.font.color.rgb = RGBColor(100, 116, 139)
            
        p_media.add_run("    ") # दुई चित्र बीचको खाली ठाउँ
        
        # १२. QR Code जेनेरेट गर्ने
        qr_text = f"ID: {player.get('id', 'N/A')} | Name: {player.get('name', '')} | School: {school_name}"
        qr = qrcode.make(qr_text)
        qr_io = BytesIO()
        qr.save(qr_io, format='PNG')
        qr_io.seek(0)
        p_media.add_run().add_picture(qr_io, height=Cm(2.0))

        # --- Details ---
        p_det = cell.add_paragraph()
        p_det.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # ११. Long Event List Truncation
        player_id = int(player['id']) if pd.notna(player.get('id')) else 0
        events_list = events_dict.get(player_id, [])
        events_str = ", ".join(events_list) if events_list else "N/A"
        if len(events_str) > 40:
            events_str = events_str[:37] + "..." # लामो भएमा छोट्याउने
            
        class_val = player.get('class_val', 'N/A')
        dob_val = player.get('dob_bs', 'N/A')
        gender_val = player.get('gender', 'N/A')
        
        details = f"नाम: {player.get('name', 'N/A')}\n"
        details += f"कक्षा: {class_val}  |  लिङ्ग: {gender_val}\n"
        details += f"जन्म मिति: {dob_val}\n"
        details += f"खेलहरू: {events_str}"
        
        r_det = p_det.add_run(details)
        r_det.font.size = Pt(10)
        r_det.font.name = 'Arial Unicode MS' # ४. नेपाली फन्ट
        r_det._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        
        # --- Signatures ---
        p_sig = cell.add_paragraph()
        p_sig.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # ६. स्पष्ट खाली ठाउँ र लाइन
        r_sig_line = p_sig.add_run("\n_______________        _______________\n")
        r_sig_line.font.size = Pt(8)
        
        r_sig_txt = p_sig.add_run("प्रअको हस्ताक्षर                 आयोजक")
        r_sig_txt.font.size = Pt(8)
        r_sig_txt.font.name = 'Arial Unicode MS'
        r_sig_txt._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')

    # १०. Streamlit को लागि BytesIO मा सेभ गरेर फर्काउने
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output