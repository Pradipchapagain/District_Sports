import pandas as pd
import database as db
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.lib import colors
from datetime import datetime



# ==========================================
# 📜 परिचय-पत्र बनाउने फङ्सन (ID CARD Generator)
# ==========================================


def generate_id_cards_docx(school_name, players_df):
    """Word फाइलमा आकर्षक परिचयपत्र (ID Cards) बनाउने फङ्सन"""
    doc = Document()
    
    # Margin Setup (Narrow margins for ID cards)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    
    # Main Header
    head = doc.add_heading('१६औं जिल्ला स्तरीय राष्ट्रपति रनिङ शिल्ड प्रतियोगिता २०८२', level=1)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head.runs[0].font.size = Pt(14)
    head.runs[0].font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
    
    sub = doc.add_paragraph(f"विद्यालय: {school_name}")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].bold = True
    
    doc.add_paragraph("") # Space
    
    # Grid Table (2 Columns)
    rows_needed = (len(players_df) + 1) // 2
    table = doc.add_table(rows=rows_needed, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set fixed column widths (approx. 8.5 cm each)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(8.5)
    
    for idx, (_, player) in enumerate(players_df.iterrows()):
        r = idx // 2
        c = idx % 2
        cell = table.cell(r, c)
        
        # Cell Content Formatting
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Card Title
        r_title = p.add_run("PLAYER IDENTITY CARD\n")
        r_title.bold = True
        r_title.font.size = Pt(12)
        r_title.font.color.rgb = RGBColor(200, 0, 0) # Red
        
        # Photo Placeholder (Box)
        r_photo = p.add_run("\n[ टाँस्ने ठाउँ ]\n(PP Size Photo)\n\n")
        r_photo.font.size = Pt(9)
        r_photo.font.color.rgb = RGBColor(128, 128, 128) # Gray
        
        # Fetch Events for the Player
        conn = db.get_connection()
        ev_query = """
            SELECT e.name 
            FROM registrations r
            JOIN events e ON r.event_code = e.code
            WHERE r.player_id = ?
        """
        events_df = pd.read_sql_query(ev_query, conn, params=(player['id'],))
        conn.close()
        
        events_list = events_df['name'].tolist() if not events_df.empty else []
        events_str = ", ".join(events_list) if events_list else "N/A"
        
        # Player Details
        details =  f"नाम (Name): {player['name']}\n"
        details += f"कक्षा (Class): {player.get('class_val', 'N/A')}\n"
        details += f"लिङ्ग (Gender): {player['gender']} | जन्म मिति: {player['dob_bs']}\n"
        details += f"सहभागी खेलहरू: {events_str}\n"
        
        r_det = p.add_run(details)
        r_det.font.size = Pt(10)
        
        # Footer / Signatures
        p.add_run("\n\n__________________         __________________\n")
        p.add_run("प्रअको हस्ताक्षर                 आयोजकको हस्ताक्षर").font.size = Pt(9)
        
    return doc

from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.lib import colors
from datetime import datetime

# ==========================================
# 📜 प्रमाणपत्र बनाउने फङ्सन (Certificate Generator)
# ==========================================
def generate_certificate_pdf(event_name, winners_df):
    """विजेताहरूको लिस्टबाट प्रमाणपत्रको PDF बनाउने"""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    
    for _, row in winners_df.iterrows():
        # डाटाबेसबाट आउने कोलमको नाम 'school_name' र 'rank' वा 'position' हुन सक्छ
        p_name = row.get('name', 'Unknown Player')
        p_school = row.get('school_name', row.get('School', 'Unknown School'))
        p_rank = int(row.get('rank', row.get('position', 1)))
        
        # 1. Borders
        c.setStrokeColor(colors.darkblue)
        c.setLineWidth(5)
        c.rect(1*cm, 1*cm, width-2*cm, height-2*cm)
        
        c.setStrokeColor(colors.gold)
        c.setLineWidth(2)
        c.rect(1.2*cm, 1.2*cm, width-2.4*cm, height-2.4*cm)
        
        # 2. Main Event Header
        c.setFont("Helvetica-Bold", 18)
        c.setFillColor(colors.darkred)
        c.drawCentredString(width/2, height - 2.5*cm, "16th District Level President Running Shield 2082")
        
        c.setFont("Helvetica-Bold", 30)
        c.setFillColor(colors.darkblue)
        c.drawCentredString(width/2, height - 4.5*cm, "CERTIFICATE OF ACHIEVEMENT")
        
        c.setFont("Helvetica", 14)
        c.setFillColor(colors.black)
        c.drawCentredString(width/2, height - 6*cm, "This certificate is proudly presented to")
        
        # 3. Player Name
        c.setFont("Helvetica-BoldOblique", 26)
        c.setFillColor(colors.darkred)
        c.drawCentredString(width/2, height - 7.8*cm, f"{p_name}")
        c.setFillColor(colors.black)
        
        # 4. School
        c.setFont("Helvetica", 16)
        c.drawCentredString(width/2, height - 9.3*cm, f"of {p_school}")
        
        # 5. Position Text
        c.setFont("Helvetica", 14)
        c.drawCentredString(width/2, height - 11.2*cm, "for securing the position of")
        
        # 6. Rank & Event
        if p_rank == 1:
            rank_text = "FIRST (Gold)"
            c.setFillColor(colors.gold)
        elif p_rank == 2:
            rank_text = "SECOND (Silver)"
            c.setFillColor(colors.silver)
        else:
            rank_text = "THIRD (Bronze)"
            c.setFillColor(colors.saddlebrown)
            
        c.setFont("Helvetica-Bold", 22)
        c.drawCentredString(width/2, height - 12.8*cm, f"{rank_text}")
        
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 16)
        c.drawCentredString(width/2, height - 14*cm, f"in the event {event_name}")
        
        # 7. Date & Signatures
        c.setFont("Helvetica", 12)
        today_date = datetime.now().strftime("%Y-%m-%d")
        
        # Left Signature (Date)
        c.line(3*cm, 3*cm, 8*cm, 3*cm)
        c.drawString(4.5*cm, 2.5*cm, "Date")
        c.drawString(4*cm, 3.2*cm, today_date)
        
        # Right Signature (Organizer)
        c.line(width-8*cm, 3*cm, width-3*cm, 3*cm)
        c.drawCentredString(width-5.5*cm, 2.5*cm, "Organizer / Coordinator")
        
        # अर्को पेज (Next Certificate)
        c.showPage()
        
    c.save()
    buffer.seek(0)
    return buffer